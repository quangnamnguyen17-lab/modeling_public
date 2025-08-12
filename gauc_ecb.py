""""libraries"""
# Import necessary libraries
import numpy as np
import pandas as pd
import seaborn as sns
import scipy as sp
import matplotlib.pyplot as plt
import random
import warnings
import inspect
import traceback
import logging
from sklearn.metrics import roc_curve, auc

# Set logging level to WARNING
logging.basicConfig(level=logging.WARNING)

from IPython.display import display
from pandas.core.arrays import numpy_
from datetime import datetime, timedelta
from scipy.stats import kruskal, spearmanr, f_oneway, ttest_ind, mannwhitneyu, ks_2samp, shapiro
from sklearn.tree import DecisionTreeClassifier, plot_tree
from sklearn.preprocessing import KBinsDiscretizer, LabelEncoder, label_binarize
from sklearn.metrics import roc_auc_score
from itertools import combinations
from typing import Dict, List, Tuple, Union
from docx import Document

def assign_level(value, intervals, levels):
    """
    Assigns a level to a value based on defined intervals.

    :param value: The value to categorize.
    :param intervals: A list of interval boundaries (e.g., [0, 10, 20, 30]).
    :param levels: A list of levels corresponding to intervals (e.g., ['Low', 'Medium', 'High']).
    :return: The level corresponding to the interval the value falls into.
    """
    for i in range(len(intervals) - 1):
        if intervals[i] <= value < intervals[i + 1]:
            return levels[i]
    return levels[-1]  # Default to the last level if value exceeds intervals.

"""PERFORMANCE GAUC"""
# Affichage du tableau de contingence
def contigence(
    y_true: pd.Series,
    y_pred: pd.Series,
    label_y_true: str="y_true",
    label_y_pred: str="y_pred"
) -> pd.DataFrame:
    """
    Affiche le tableau de contingence et sa heatmap.
    """
    cont_tab = pd.crosstab(y_true, y_pred)

    # Visualisation avec un heatmap
    plt.figure(figsize=(8, 6))
    sns.heatmap(cont_tab, annot=True, fmt='d', cmap='Blues')
    plt.title("Tableau de contingence (Heatmap)")
    plt.xlabel(label_y_pred)
    plt.ylabel(label_y_true)
    plt.show()
    return cont_tab


def generalized_auc_ecb(
    y_true: pd.Series,
    y_pred: pd.Series,
    min_threshold: float = 0.6,
    warning_threshold: float = 0.5,
    method: str = "quantile",
    verbose: bool = True
) -> Tuple[float, str]:
    """
    Calcule un AUC généralisé entre y_true (continue) et y_pred (catégorielle),
    après transformation ordonnée et alignement des classes.

    Paramètres :
    ----------
    y_true : pd.Series
        Variable continue (ex: LGD_OBS).
    y_pred : pd.Series
        Variable catégorielle prédite (ex: segment).
    min_threshold : float
        Seuil "orange".
    warning_threshold : float
        Seuil "rouge".
    verbose : bool
        Affiche les étapes de transformation si True.

    Retour :
    -------
    Tuple[float, str] : AUC généralisé, Couleur associée ("vert", "orange", "rouge").
    """

    df = pd.DataFrame({"y_true": y_true, "y_pred": y_pred}).dropna()

    bins = [-np.inf,0.05,0.1,0.2,0.3,0.4,0.5,0.6,0.7,0.8,0.9,0.95,np.inf]    # Discrétisation de y_true

    df["y_true_binned"] = pd.cut(df["y_true"],bins =bins,include_lowest = True)
    df["y_true_binned"].value_counts()
    
    # Encodage ordonné de y_pred selon moyenne de LGD
    LGD_pred_dict=dict()
    LGD_pred_dict = df.groupby(['y_pred'])['y_true'].mean().to_dict()

    df["y_pred_mean"] = df["y_pred"].map(LGD_pred_dict)
    
    df["y_pred_binned"] = pd.cut(df["y_pred_mean"],bins =bins,include_lowest = True)
    df["y_pred_binned"].value_counts()

    if verbose:
        print("\nPrésentation des classes de y_true après discretisation :")
        print(df["y_true_binned"].value_counts().sort_index())

    # Label binarize les deux pour AUC multiclass
    y_true_bin = df["y_true_binned"]
    y_pred_bin = df["y_pred_binned"]
    

    tabcont=contigence(df["y_true_binned"],df["y_pred_binned"])

    somersdyx = sp.stats.somersd(df["y_pred_binned"], df["y_true_binned"], alternative='two-sided')
    somersdxy = sp.stats.somersd(df["y_true_binned"], df["y_pred_binned"], alternative='two-sided')
    print("Somers' D :", somersdyx.statistic, ", pvalue : ", somersdyx.pvalue)
    print("Somers' D :", somersdxy.statistic, ", pvalue : ", somersdxy.pvalue)
    

    auc = (somersdyx.statistic + 1) /2
    print("GAUC : ", auc)


    levels = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12]  # Define corresponding levels

    y_true_encoded = [assign_level(v, bins, levels) for v in df["y_true"]]

    y_pred_encoded = [assign_level(v, bins, levels) for v in df["y_pred_mean"]]   

    """print("y_true_encoded :", y_true_encoded)
    print("y_pred_encoded :", y_pred_encoded)
    
    auc = roc_auc_score(y_true_encoded,y_pred_encoded, multi_class="ovr")
    cumulative_roc(df["y_true"], df["y_pred_mean"])"""

    # Interprétation
    if auc < warning_threshold:
        color = "\U0001F534" # rouge
    elif auc < min_threshold:
        color = "\U0001F7E0" # orange
    else:
        color = "\U0001F7E2" # vert

    return auc, color


def export_doc(functions: list, file_name: str = "documentation.docx"):
    doc = Document()
    doc.add_heading("Documentation des fonctions de backtesting de LGD", 0)

    for func in functions:
        doc.add_heading(func.__name__, level=1)
        doc.add_paragraph(inspect.getdoc(func) or "Pas de docstring")
        doc.add_paragraph("Signature :")
        doc.add_paragraph(str(inspect.signature(func)))

    doc.save(file_name)

"""
export_doc(functions=[interpret, run_group_comparison_tests, generalized_auc, generalized_auc_ecb], file_name="backtesting_lgd_documentation_QNN.docx")
"""
